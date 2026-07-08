import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def create_report(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = docx.Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    in_code_block = False
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            if not in_code_block:
                continue
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                continue
        
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        if stripped == '---':
            continue
            
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, stripped[2:])
        elif stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, stripped[stripped.index('.')+2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)
            
    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    create_report("report_final.md", "CST209_Final_Report.docx")
